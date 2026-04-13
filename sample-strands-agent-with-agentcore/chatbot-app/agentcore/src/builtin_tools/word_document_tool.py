"""
Word Document Tools - 5 essential tools for Word document management.

Tools:
1. create_word_document - Create new Word document from Python code
2. modify_word_document - Modify existing Word document with python-docx code
3. list_my_word_documents - List all Word documents in workspace
4. read_word_document - Retrieve document for download
5. preview_word_page - Get page screenshot for visual inspection

Note: Uploaded .docx files are automatically stored to workspace by agent.py
Pattern follows diagram_tool for Code Interpreter usage.
"""

import os
import re
import logging
from typing import Dict, Any, Optional
from strands import tool, ToolContext
from skill import register_skill
from workspace import WordManager
from builtin_tools.lib.tool_response import build_success_response, build_image_response

logger = logging.getLogger(__name__)


def _validate_document_name(name: str) -> tuple[bool, Optional[str]]:
    """Validate document name meets requirements (without extension).

    Rules:
    - Only letters (a-z, A-Z), numbers (0-9), hyphens (-), and underscores (_)
    - No spaces or special characters
    - No consecutive hyphens
    - No leading/trailing hyphens

    Args:
        name: Document name without extension (e.g., "sales-report")

    Returns:
        (is_valid, error_message)
        - (True, None) if valid
        - (False, error_message) if invalid
    """
    # Check for empty name
    if not name:
        return False, "Document name cannot be empty"

    # Check for valid characters: letters, numbers, hyphens, underscores
    if not re.match(r'^[a-zA-Z0-9_\-]+$', name):
        invalid_chars = re.findall(r'[^a-zA-Z0-9_\-]', name)
        return False, f"Invalid characters in name: {set(invalid_chars)}. Use only letters, numbers, hyphens, and underscores."

    # Check for consecutive hyphens
    if '--' in name:
        return False, "Name cannot contain consecutive hyphens (--)"

    # Check for leading/trailing hyphens
    if name.startswith('-') or name.endswith('-'):
        return False, "Name cannot start or end with a hyphen"

    return True, None


def _sanitize_document_name_for_bedrock(filename: str) -> str:
    """Sanitize existing filename for Bedrock API (removes extension).

    Use this ONLY for existing files being read from S3.
    For new files, use _validate_document_name() instead.

    Args:
        filename: Original filename with extension (e.g., "test_document_v2.docx")

    Returns:
        Sanitized name without extension (e.g., "test-document-v2")
    """
    # Remove extension
    if '.' in filename:
        name, ext = filename.rsplit('.', 1)
    else:
        name = filename

    # Replace underscores and spaces with hyphens
    name = name.replace('_', '-').replace(' ', '-')

    # Keep only allowed characters: alphanumeric, hyphens, parentheses, square brackets
    # This matches agent.py's _sanitize_filename behavior
    name = re.sub(r'[^a-zA-Z0-9\-\(\)\[\]]', '', name)

    # Replace multiple consecutive hyphens with single hyphen
    name = re.sub(r'\-+', '-', name)

    # Trim hyphens from start/end
    name = name.strip('-')

    # If name becomes empty, use default
    if not name:
        name = 'document'

    if name != filename.replace('.docx', ''):
        logger.info(f"Sanitized document name for Bedrock: '{filename}' → '{name}'")

    return name


def _get_user_session_ids(tool_context: ToolContext) -> tuple[str, str]:
    """Extract user_id and session_id from ToolContext

    Returns:
        (user_id, session_id) tuple
    """
    # Extract from invocation_state (set by agent or swarm)
    invocation_state = tool_context.invocation_state
    user_id = invocation_state.get('user_id', 'default_user')
    session_id = invocation_state.get('session_id', 'default_session')

    logger.info(f"Extracted IDs: user_id={user_id}, session_id={session_id}")
    return user_id, session_id


def _save_word_artifact(
    tool_context: ToolContext,
    filename: str,
    s3_url: str,
    size_kb: str,
    tool_name: str,
    user_id: str,
    session_id: str
) -> None:
    """Save Word document as artifact to agent.state for Canvas display.

    Args:
        tool_context: Strands ToolContext
        filename: Document filename (e.g., "report.docx")
        s3_url: Full S3 URL (e.g., "s3://bucket/path/report.docx")
        size_kb: File size string (e.g., "45.2 KB")
        tool_name: Tool that created this ("create_word_document" or "modify_word_document")
        user_id: User ID
        session_id: Session ID
    """
    from datetime import datetime, timezone

    try:
        # Generate artifact ID using filename (without extension) for easy lookup
        doc_name = filename.replace('.docx', '')
        artifact_id = f"word-{doc_name}"

        # Get current artifacts from agent.state
        artifacts = tool_context.agent.state.get("artifacts") or {}

        # Create/update artifact
        artifacts[artifact_id] = {
            "id": artifact_id,
            "type": "word_document",
            "title": filename,
            "content": s3_url,  # Full S3 URL for OfficeViewer
            "tool_name": tool_name,
            "metadata": {
                "filename": filename,
                "s3_url": s3_url,
                "size_kb": size_kb,
                "user_id": user_id,
                "session_id": session_id
            },
            "created_at": artifacts.get(artifact_id, {}).get("created_at", datetime.now(timezone.utc).isoformat()),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }

        # Save to agent.state
        tool_context.agent.state.set("artifacts", artifacts)

        # Sync agent state to persistence
        session_manager = tool_context.invocation_state.get("session_manager")
        if not session_manager and hasattr(tool_context.agent, 'session_manager'):
            session_manager = tool_context.agent.session_manager

        if session_manager:
            session_manager.sync_agent(tool_context.agent)
            logger.info(f"Saved Word artifact: {artifact_id}")
        else:
            logger.warning(f"No session_manager found, Word artifact not persisted: {artifact_id}")

    except Exception as e:
        logger.error(f"Failed to save Word artifact: {e}")


@tool(context=True)
def create_word_document(
    python_code: str,
    document_name: str,
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Create a new Word document using python-docx code.

    This tool executes python-docx code to create a document from scratch.
    Perfect for generating structured documents with headings, paragraphs, tables, and charts.

    Available libraries: python-docx, matplotlib, pandas, numpy

    Args:
        python_code: Python code using python-docx to build the document.
                    The document is initialized as: doc = Document()
                    After your code, it's automatically saved.

                    DO NOT include Document() initialization or doc.save() calls.

                    Uploaded images are automatically available in Code Interpreter.
                    Use os.listdir() to discover available image files.

                    Common Patterns:

                    Basic Structure:
                    ```python
doc.add_heading('Quarterly Report', level=1)
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph('Revenue increased by 15%...')

# Table with data
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Quarter'
table.rows[0].cells[1].text = 'Revenue'
                    ```

                    With Generated Chart:
                    ```python
import matplotlib.pyplot as plt
from docx.shared import Inches

doc.add_heading('Sales Analysis', level=1)

# Generate chart
plt.figure(figsize=(8, 5))
plt.bar(['Q1','Q2','Q3','Q4'], [100, 120, 150, 140])
plt.title('Quarterly Sales')
plt.savefig('sales.png', dpi=300, bbox_inches='tight')
plt.close()

# Insert chart
doc.add_paragraph().add_run().add_picture('sales.png', width=Inches(6))
doc.add_paragraph('Figure 1: Sales performance')
                    ```

                    With Uploaded Image:
                    ```python
from docx.shared import Inches
import os

doc.add_heading('Product Catalog', level=1)
doc.add_paragraph('Our new product line:')

# Discover available images
available_images = [f for f in os.listdir() if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp'))]

if available_images:
    # Use the first image (or select specific one by filename matching)
    image_file = available_images[0]
    doc.add_paragraph().add_run().add_picture(image_file, width=Inches(5))
    doc.add_paragraph(f'Figure: {image_file}')
else:
    # No images found - add placeholder
    doc.add_paragraph('[No images available in workspace]')
                    ```

                    With Hyperlinks:
                    ```python
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

para = doc.add_paragraph('Visit ')
add_hyperlink(para, 'our website', 'https://example.com')
                    ```

                    With Placeholders for Future Edits:
                    ```python
doc.add_heading('Analysis', level=2)
doc.add_paragraph('{{INSERT_CHART_HERE}}')  # Marker for modify_word_document
doc.add_paragraph('Summary text...')
                    ```

        document_name: Document name WITHOUT extension (.docx is added automatically)
                      Use ONLY letters, numbers, hyphens (no underscores or spaces)
                      Examples: "sales-report", "Q4-analysis", "report-final"

    Returns:
        Success message with file details and workspace list

    Note:
        - Document is saved to workspace for future editing with modify_word_document
        - Uploaded images are automatically available in Code Interpreter
        - Keep code focused on structure; use modify_word_document for complex refinements
    """
    try:
        logger.info("=== create_word_document called ===")
        logger.info(f"Document name: {document_name}")

        # Validate document name (without extension)
        is_valid, error_msg = _validate_document_name(document_name)
        if not is_valid:
            return {
                "content": [{
                    "text": f"**Invalid document name**: {document_name}\n\n{error_msg}\n\n**Examples of valid names:**\n- sales-report\n- Q4-analysis\n- report-final-v2"
                }],
                "status": "error"
            }

        # Add .docx extension
        document_filename = f"{document_name}.docx"
        logger.info(f"Full filename: {document_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordManager(user_id, session_id)

        # Get shared CI client (persistent across calls — never stop it here)
        from builtin_tools.code_interpreter_tool import get_ci_session
        code_interpreter = get_ci_session(tool_context)
        if code_interpreter is None:
            return {
                "content": [{
                    "text": "**Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        # Load all workspace images from S3 to Code Interpreter
        loaded_images = doc_manager.load_workspace_images_to_ci(code_interpreter)
        if loaded_images:
            logger.info(f"Loaded {len(loaded_images)} image(s) from workspace: {loaded_images}")

        # Get Code Interpreter path for file (filename only, no subdirectory)
        ci_path = doc_manager.get_ci_path(document_filename)

        # Build document creation code
        creation_code = f"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Execute user's creation code
{python_code}

# Save document
doc.save('{ci_path}')
print(f"Document created: {ci_path}")
"""

        # Execute creation
        response = code_interpreter.invoke("executeCode", {
            "code": creation_code,
            "language": "python",
            "clearContext": False
        })

        # Capture stdout and check for errors
        stdout_output = ""
        for event in response.get("stream", []):
            result = event.get("result", {})
            if result.get("isError", False):
                error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                logger.error(f"Creation failed: {error_msg[:500]}")
                return {
                    "content": [{
                        "text": f"**Failed to create document**\n\n```\n{error_msg[:1000]}\n```\n\nTip:Check your python-docx code for syntax errors or incorrect API usage."
                    }],
                    "status": "error"
                }
            # Capture stdout
            stdout = result.get("structuredContent", {}).get("stdout", "")
            if stdout:
                stdout_output += stdout

        logger.info("Document creation completed")

        # Download from Code Interpreter
        file_bytes = doc_manager.download_from_code_interpreter(code_interpreter, document_filename)

        # Save to S3 for persistence
        s3_info = doc_manager.save_to_s3(
            document_filename,
            file_bytes,
            metadata={'source': 'python_code_creation'}
        )

        # Save as artifact for Canvas display
        _save_word_artifact(
            tool_context=tool_context,
            filename=document_filename,
            s3_url=s3_info['s3_url'],
            size_kb=s3_info['size_kb'],
            tool_name='create_word_document',
            user_id=user_id,
            session_id=session_id
        )

        # Get current workspace list
        workspace_docs = doc_manager.list_s3_documents()
        other_files_count = len([d for d in workspace_docs if d['filename'] != document_filename])

        message = f"""**Document created successfully**

**File**: {document_filename} ({s3_info['size_kb']})
**Other files in workspace**: {other_files_count} document{'s' if other_files_count != 1 else ''}"""

        # Include stdout output if any
        if stdout_output.strip():
            message += f"\n\n**Output:**\n```\n{stdout_output.strip()}\n```"

        # Return success message
        return build_success_response(message, {
            "filename": document_filename,
            "tool_type": "word_document",
            "user_id": user_id,
            "session_id": session_id
        })

    except Exception as e:
        logger.error(f"create_word_document failed: {e}")
        return {
            "content": [{
                "text": f"**Failed to create document**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def modify_word_document(
    source_name: str,
    output_name: str,
    python_code: str,
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Modify existing Word document using python-docx code and save with a new name.

    This tool loads a document from workspace, executes python-docx code to modify it,
    and saves it with a new filename to preserve the original.

    Available libraries: python-docx, matplotlib, pandas, numpy

    IMPORTANT Safety Rules:
    - Always use different output_filename than source_filename (e.g., "report.docx" → "report_v2.docx")
    - Always check array lengths before accessing (len(doc.paragraphs))
    - Use try-except for operations that might fail

    Args:
        source_name: Document name to load (WITHOUT extension, must exist in workspace)
                    Example: "sales-report", "Q4-analysis"
        output_name: New document name (WITHOUT extension, must be different from source)
                    Use ONLY letters, numbers, hyphens (no underscores or spaces)
                    Example: "sales-report-v2", "Q4-analysis-final"
        python_code: Python code using python-docx library to modify document.
                    The document is loaded as: doc = Document('<filename>')
                    After modifications, it's automatically saved.

                    DO NOT include Document() initialization or doc.save() calls.

                    Uploaded images are automatically available in Code Interpreter.
                    Use os.listdir() to discover available image files.

                    Common Patterns:

                    Insert Chart at Marker:
                    ```python
import matplotlib.pyplot as plt
from docx.shared import Inches

# Generate chart
plt.figure(figsize=(8, 5))
plt.plot([1,2,3,4], [10, 20, 25, 30])
plt.title('Sales Trend')
plt.savefig('trend.png', dpi=300, bbox_inches='tight')
plt.close()

# Find marker and replace with chart
for para in doc.paragraphs:
    if '{{CHART}}' in para.text:
        para.clear()
        para.add_run().add_picture('trend.png', width=Inches(6))
        break
                    ```

                    Insert Uploaded Image:
                    ```python
from docx.shared import Inches
import os

# Discover available images
available_images = [f for f in os.listdir() if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp'))]

if available_images:
    # Use the first image (or select specific one by filename matching)
    image_file = available_images[0]

    # Add new paragraph with image
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(image_file, width=Inches(6.5))

    # Add caption
    caption = doc.add_paragraph(f'Figure: {image_file}')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    # No images found - add placeholder text
    doc.add_paragraph('[Image placeholder - no images found in workspace]')
                    ```

                    Add Hyperlink:
                    ```python
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Add to end of document
para = doc.add_paragraph('For more info: ')
add_hyperlink(para, 'Click here', 'https://example.com')
                    ```

                    Preserve Formatting When Editing:
                    ```python
# Preserve existing formatting when modifying text
if len(doc.paragraphs) > 0:
    p = doc.paragraphs[0]
    if len(p.runs) > 0:
        # Copy original formatting
        original_run = p.runs[0]
        font_name = original_run.font.name
        font_size = original_run.font.size
        is_bold = original_run.font.bold

        # Clear and add new text with same formatting
        for run in p.runs:
            run.text = ''

        new_run = p.runs[0] if len(p.runs) > 0 else p.add_run()
        new_run.text = 'New text with preserved formatting'
        new_run.font.name = font_name
        new_run.font.size = font_size
        new_run.font.bold = is_bold
                    ```

    Returns:
        Success message with file details and workspace list

    Note:
        - Uploaded images are automatically available in Code Interpreter
        - Use 0-based indexing (first paragraph = index 0)
        - Document automatically synced to S3
    """
    try:
        logger.info("=== modify_word_document called ===")
        logger.info(f"Source: {source_name}, Output: {output_name}")

        # Validate output name format
        is_valid, error_msg = _validate_document_name(output_name)
        if not is_valid:
            return {
                "content": [{
                    "text": f"**Invalid output name**: {output_name}\n\n{error_msg}\n\n**Examples of valid names:**\n- sales-report-v2\n- Q4-analysis-final\n- report-revised"
                }],
                "status": "error"
            }

        # Ensure source and output are different
        if source_name == output_name:
            return {
                "content": [{
                    "text": f"**Invalid name**\n\nOutput name must be different from source name to preserve the original.\n\nSource: {source_name}\nOutput: {output_name}\n\nTip:Try: \"{source_name}-v2\""
                }],
                "status": "error"
            }

        # Add .docx extensions
        source_filename = f"{source_name}.docx"
        output_filename = f"{output_name}.docx"
        logger.info(f"Full filenames: {source_filename} → {output_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordManager(user_id, session_id)

        # Get shared CI client (persistent across calls — never stop it here)
        from builtin_tools.code_interpreter_tool import get_ci_session
        code_interpreter = get_ci_session(tool_context)
        if code_interpreter is None:
            return {
                "content": [{
                    "text": "**Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        # Load all workspace images from S3 to Code Interpreter
        loaded_images = doc_manager.load_workspace_images_to_ci(code_interpreter)
        if loaded_images:
            logger.info(f"Loaded {len(loaded_images)} image(s) from workspace: {loaded_images}")

        # Ensure source file is in Code Interpreter (load from S3 if needed)
        source_ci_path = doc_manager.ensure_file_in_ci(code_interpreter, source_filename)

        # Generate output path
        output_ci_path = doc_manager.get_ci_path(output_filename)

        # Build modification code
        modification_code = f"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load source document
doc = Document('{source_ci_path}')

# Execute user's modification code
{python_code}

# Save to output file
doc.save('{output_ci_path}')
print(f"Document modified and saved: {output_ci_path}")
"""

        # Execute modification
        response = code_interpreter.invoke("executeCode", {
            "code": modification_code,
            "language": "python",
            "clearContext": False
        })

        # Capture stdout and check for errors
        stdout_output = ""
        for event in response.get("stream", []):
            result = event.get("result", {})
            if result.get("isError", False):
                error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                logger.error(f"Modification failed: {error_msg[:500]}")
                return {
                    "content": [{
                        "text": f"**Modification failed**\n\n```\n{error_msg[:1000]}\n```\n\nTip:Check your python-docx code for syntax errors or incorrect API usage."
                    }],
                    "status": "error"
                }
            # Capture stdout
            stdout = result.get("structuredContent", {}).get("stdout", "")
            if stdout:
                stdout_output += stdout

        logger.info("Document modification completed")

        # Download modified document from Code Interpreter
        file_bytes = doc_manager.download_from_code_interpreter(code_interpreter, output_filename)

        # Save to S3 with output filename
        s3_info = doc_manager.save_to_s3(
            output_filename,
            file_bytes,
            metadata={
                'source': 'modification',
                'source_filename': source_filename,
                'modified_at': 'timestamp'
            }
        )

        # Save as artifact for Canvas display
        _save_word_artifact(
            tool_context=tool_context,
            filename=output_filename,
            s3_url=s3_info['s3_url'],
            size_kb=s3_info['size_kb'],
            tool_name='modify_word_document',
            user_id=user_id,
            session_id=session_id
        )

        # Get current workspace list
        workspace_docs = doc_manager.list_s3_documents()
        other_files_count = len([d for d in workspace_docs if d['filename'] != output_filename])

        # Build success message
        message = f"""**Document modified successfully**

**Source**: {source_filename}
**Saved as**: {output_filename} ({s3_info['size_kb']})
**Other files in workspace**: {other_files_count} document{'s' if other_files_count != 1 else ''}"""

        # Include stdout output if any
        if stdout_output.strip():
            message += f"\n\n**Output:**\n```\n{stdout_output.strip()}\n```"

        # Return success message with metadata for download button
        return build_success_response(message, {
            "filename": output_filename,
            "tool_type": "word_document",
            "user_id": user_id,
            "session_id": session_id
        })

    except FileNotFoundError as e:
        logger.error(f"Document not found: {e}")
        return {
            "content": [{
                "text": f"**Document not found**: {source_filename}"
            }],
            "status": "error"
        }
    except Exception as e:
        logger.error(f"modify_word_document failed: {e}")
        return {
            "content": [{
                "text": f"**Failed to modify document**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def list_my_word_documents(
    tool_context: ToolContext
) -> Dict[str, Any]:
    """List all Word documents in workspace.

    Shows all .docx files in workspace with size and metadata.

    Use this tool when:
    - User asks "what Word files do I have?"
    - User says "show my documents", "list files"
    - Before modifying: verify document exists
    - User wants to see workspace contents

    No arguments needed.

    Returns:
        - Formatted list of all Word documents
        - Each entry shows: filename, size, last modified date
        - Total file count
        - Metadata for frontend download buttons

    Example Usage:
        Scenario 1 - Check available files:
            User: "What Word documents do I have?"
            AI: list_my_word_documents()
            → Shows: report.docx, proposal.docx, analysis.docx

        Scenario 2 - Before modifying:
            User: "Edit my report"
            AI: [Unclear which file]
            AI: list_my_word_documents()
            AI: "I found these documents: ... Which one should I modify?"

    Example Output:
        Workspace (3 documents):
          - q4_report.docx (45.6 KB) - Modified: 2025-01-15
          - proposal.docx (32.1 KB) - Modified: 2025-01-14
          - analysis.docx (78.4 KB) - Modified: 2025-01-13

    Note:
        - Shows files from workspace
        - Empty workspace shows helpful message
        - Frontend renders download buttons automatically
    """
    try:
        logger.info("=== list_my_word_documents called ===")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordManager(user_id, session_id)

        # List documents from S3
        documents = doc_manager.list_s3_documents()

        # Format list
        workspace_summary = doc_manager.format_file_list(documents)

        if documents:
            message = workspace_summary
        else:
            message = workspace_summary

        # Prepare metadata for frontend (download buttons)
        metadata = {
            "documents": [
                {
                    "filename": doc['filename'],
                    "s3_key": doc['s3_key'],
                    "size_kb": doc['size_kb'],
                    "last_modified": doc['last_modified']
                } for doc in documents
            ]
        }

        return build_success_response(message, metadata)

    except Exception as e:
        logger.error(f"list_my_word_documents failed: {e}")
        return {
            "content": [{
                "text": f"**Failed to list documents**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def read_word_document(
    document_name: str,
    tool_context: ToolContext,
    include_comments: bool = False
) -> Dict[str, Any]:
    """Read and retrieve a specific Word document.

    This tool loads a document from workspace and extracts its text content using Code Interpreter.
    The extracted text (paragraphs, tables, etc.) is returned for analysis and answering questions.

    Use this tool when:
    - User asks about document contents: "What's in report.docx?", "Summarize this document"
    - User wants to analyze the document: "How many tables are in this file?", "What's the main topic?"
    - User explicitly requests download: "Send me [filename]", "I need [document]"
    - You need to verify document contents before modification
    - User asks about comments/reviews: "What comments are in this doc?", "Show me the review feedback"

    IMPORTANT:
    - For creating new documents: use create_word_document
    - For modifying documents: use modify_word_document

    Args:
        document_name: Document name WITHOUT extension (.docx is added automatically)
                      Must exist in workspace.
                      Example: "report", "proposal", "Q4-analysis"
        include_comments: If True, extract and display comments with their locations.
                         Each comment shows: author, date, text, and which paragraph it's attached to.
                         Default: False

    Returns:
        - Extracted text content (paragraphs, tables, headings)
        - If include_comments=True: comments with paragraph mapping
        - Document metadata (filename, size, S3 location)
        - Frontend shows download button based on metadata

    Example Usage:
        # Download request
        User: "Send me the report"
        AI: read_word_document("report.docx")

        # Read with comments
        User: "Show me the comments in the report"
        AI: read_word_document("report", include_comments=True)

    Note:
        - File must exist in workspace
        - Frontend handles download automatically
    """
    try:
        logger.info("=== read_word_document called ===")
        logger.info(f"Document name: {document_name}")

        # Add .docx extension
        document_filename = f"{document_name}.docx"
        logger.info(f"Full filename: {document_filename}")

        # Get user and session IDs
        user_id, session_id = _get_user_session_ids(tool_context)

        # Initialize document manager
        doc_manager = WordManager(user_id, session_id)

        # Load from S3
        file_bytes = doc_manager.load_from_s3(document_filename)

        # Get file info
        documents = doc_manager.list_s3_documents()
        doc_info = next((d for d in documents if d['filename'] == document_filename), None)

        if not doc_info:
            raise FileNotFoundError(f"Document not found: {document_filename}")

        # Get shared CI session
        from builtin_tools.code_interpreter_tool import get_ci_session
        code_interpreter = get_ci_session(tool_context)
        if code_interpreter is None:
            return {
                "content": [{
                    "text": "**Code Interpreter not configured**\n\nCODE_INTERPRETER_ID not found in environment or Parameter Store."
                }],
                "status": "error"
            }

        try:
            # Upload document to Code Interpreter
            doc_manager.upload_to_code_interpreter(code_interpreter, document_filename, file_bytes)

            # Generate extraction code
            comments_block = ""
            if include_comments:
                comments_block = f'''
# Extract comments via zipfile (more reliable than python-docx rels API)
import zipfile
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_NS = '{{}}'.format(W)

comments_map = {{}}
try:
    with zipfile.ZipFile("{document_filename}", 'r') as z:
        if 'word/comments.xml' in z.namelist():
            with z.open('word/comments.xml') as f:
                croot = etree.parse(f).getroot()
            for c in croot.findall('.//{{' + W + '}}comment'):
                cid = c.get('{{' + W + '}}id')
                author = c.get('{{' + W + '}}author', '')
                date = c.get('{{' + W + '}}date', '')
                texts = []
                for t in c.findall('.//{{' + W + '}}t'):
                    if t.text:
                        texts.append(t.text)
                comments_map[cid] = {{
                    "id": cid,
                    "author": author,
                    "date": date,
                    "text": ''.join(texts),
                }}
except Exception as e:
    pass

# Map comment IDs to paragraph indices via commentRangeStart in document body
comment_to_para = {{}}
body = doc.element.body
W_TAG = '{{' + W + '}}'
para_elements = body.findall(W_TAG + 'p')
for para_idx, p_elem in enumerate(para_elements):
    for marker in p_elem.iter(W_TAG + 'commentRangeStart'):
        cid = marker.get(W_TAG + 'id')
        if cid:
            comment_to_para[cid] = para_idx

# Build comments list with paragraph context
result["comments"] = []
for cid, info in comments_map.items():
    entry = dict(info)
    entry["paragraph_index"] = comment_to_para.get(cid)
    pidx = comment_to_para.get(cid)
    if pidx is not None and pidx < len(doc.paragraphs):
        ptext = doc.paragraphs[pidx].text
        entry["paragraph_text"] = ptext[:100] + ("..." if len(ptext) > 100 else "")
    result["comments"].append(entry)
result["properties"]["comments_count"] = len(comments_map)
'''

            extraction_code = f'''
import json
from docx import Document

doc = Document("{document_filename}")
result = {{
    "paragraphs": [],
    "tables": [],
    "sections": []
}}

# Extract paragraphs with styles
for para in doc.paragraphs:
    if para.text.strip():
        result["paragraphs"].append({{
            "text": para.text,
            "style": para.style.name if para.style else "Normal"
        }})

# Extract tables
for table_idx, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    result["tables"].append({{
        "index": table_idx,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "data": table_data
    }})

# Document properties
result["properties"] = {{
    "sections": len(doc.sections),
    "paragraphs_count": len(doc.paragraphs),
    "tables_count": len(doc.tables)
}}

{comments_block}

print(json.dumps(result, ensure_ascii=False))
'''

            # Execute extraction
            response = code_interpreter.invoke("executeCode", {
                "code": extraction_code,
                "language": "python",
                "clearContext": False
            })

            # Collect JSON output
            json_output = ""
            for event in response.get("stream", []):
                result = event.get("result", {})
                if result.get("isError", False):
                    error_msg = result.get("structuredContent", {}).get("stderr", "Unknown error")
                    logger.error(f"Extraction failed: {error_msg[:500]}")
                    return {
                        "content": [{
                            "text": f"**Failed to read document**\n\n```\n{error_msg[:1000]}\n```"
                        }],
                        "status": "error"
                    }

                stdout = result.get("structuredContent", {}).get("stdout", "")
                if stdout:
                    json_output += stdout

            # Parse JSON result
            import json
            doc_content = json.loads(json_output)

            # Format output text
            output_parts = []
            output_parts.append(f"📄 **Document Content**: {document_filename} ({doc_info['size_kb']})")
            output_parts.append("")

            # Format paragraphs by style
            current_style = None
            for para in doc_content.get("paragraphs", []):
                style = para.get("style", "Normal")
                text = para.get("text", "")

                if "Heading 1" in style:
                    output_parts.append(f"# {text}")
                elif "Heading 2" in style:
                    output_parts.append(f"## {text}")
                elif "Heading 3" in style:
                    output_parts.append(f"### {text}")
                elif "Title" in style:
                    output_parts.append(f"**{text}**")
                else:
                    output_parts.append(text)
                output_parts.append("")

            # Format tables
            for table in doc_content.get("tables", []):
                output_parts.append(f"**[Table {table['index'] + 1}]** ({table['rows']} rows × {table['cols']} cols)")
                for row_idx, row in enumerate(table.get("data", [])):
                    output_parts.append(" | ".join(row))
                    if row_idx == 0:
                        output_parts.append(" | ".join(["---"] * len(row)))
                output_parts.append("")

            # Format comments if present
            comments = doc_content.get("comments", [])
            if comments:
                output_parts.append("---")
                output_parts.append(f"**Comments ({len(comments)})**")
                output_parts.append("")
                for c in comments:
                    para_idx = c.get("paragraph_index")
                    para_text = c.get("paragraph_text", "")
                    location = f"paragraph {para_idx}" if para_idx is not None else "unlinked"
                    date_str = c.get("date", "")[:10] if c.get("date") else ""
                    output_parts.append(f"- **[{c.get('id')}]** {c.get('author', 'Unknown')}{' (' + date_str + ')' if date_str else ''}: {c.get('text', '')}")
                    if para_text:
                        output_parts.append(f"  > *on {location}*: \"{para_text}\"")
                    output_parts.append("")

            # Add summary
            props = doc_content.get("properties", {})
            summary_parts = [
                f"{props.get('paragraphs_count', 0)} paragraphs",
                f"{props.get('tables_count', 0)} tables",
            ]
            if props.get("comments_count"):
                summary_parts.append(f"{props['comments_count']} comments")
            output_parts.append(f"---\n*{', '.join(summary_parts)}*")
            output_parts.append(f"*Last Modified: {doc_info['last_modified'].split('T')[0]}*")

            output_text = "\n".join(output_parts)

            # Truncate if too long
            max_chars = 15000
            if len(output_text) > max_chars:
                output_text = output_text[:max_chars] + f"\n\n... (truncated, total {len(output_text)} characters)"

            return build_success_response(output_text, {
                "filename": document_filename,
                "s3_key": doc_manager.get_s3_key(document_filename),
                "size_kb": doc_info['size_kb'],
                "last_modified": doc_info['last_modified'],
                "tool_type": "word_document",
                "user_id": user_id,
                "session_id": session_id
            })

        except Exception as e:
            logger.error(f"CI execution error in read_word_document: {e}")
            raise

    except FileNotFoundError as e:
        logger.error(f"Document not found: {e}")
        return {
            "content": [{
                "text": f"**Document not found**: {document_filename}"
            }],
            "status": "error"
        }
    except Exception as e:
        logger.error(f"read_word_document failed: {e}")
        return {
            "content": [{
                "text": f"**Failed to read document**\n\n{str(e)}"
            }],
            "status": "error"
        }


@tool(context=True)
def preview_word_page(
    document_name: str,
    page_numbers: list[int],
    tool_context: ToolContext
) -> Dict[str, Any]:
    """Get page screenshots for YOU (the agent) to visually inspect before editing.

    This tool is for YOUR internal use - to see the actual layout, formatting,
    and content of pages before making modifications. Images are sent to you,
    not displayed to the user.

    Args:
        document_name: Document name without extension (e.g., "report")
        page_numbers: List of page numbers to preview (1-indexed, e.g., [1, 2, 3])

    Use BEFORE modifying a document to:
    - See exact text positions and formatting
    - Identify tables, images, or charts on pages
    - Plan precise edits based on visual layout
    """
    import subprocess
    import tempfile
    import base64
    from pdf2image import convert_from_path

    # Get user and session IDs
    user_id, session_id = _get_user_session_ids(tool_context)

    # Validate page numbers
    if not page_numbers:
        return {
            "content": [{"text": "At least one page number is required"}],
            "status": "error"
        }

    if any(p < 1 for p in page_numbers):
        return {
            "content": [{"text": "All page numbers must be 1 or greater"}],
            "status": "error"
        }

    # Validate and prepare filename
    is_valid, error_msg = _validate_document_name(document_name)
    if not is_valid:
        # Try sanitizing existing filename
        document_name = _sanitize_document_name_for_bedrock(document_name)

    document_filename = f"{document_name}.docx"
    logger.info(f"preview_word_page: {document_filename}, pages {page_numbers}")

    try:
        # Initialize document manager
        doc_manager = WordManager(user_id, session_id)

        # Check if document exists
        documents = doc_manager.list_s3_documents()
        doc_info = next((d for d in documents if d['filename'] == document_filename), None)

        if not doc_info:
            available = [d['filename'] for d in documents]
            return {
                "content": [{
                    "text": f"**Document not found**: {document_filename}\n\n"
                           f"Available documents: {', '.join(available) if available else 'None'}"
                }],
                "status": "error"
            }

        # Download Word document from S3
        docx_bytes = doc_manager.load_from_s3(document_filename)

        with tempfile.TemporaryDirectory() as temp_dir:
            # Save Word document to temp file
            docx_path = os.path.join(temp_dir, document_filename)
            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)

            # Convert Word to PDF using LibreOffice
            logger.info(f"Converting {document_filename} to PDF...")
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, docx_path],
                capture_output=True,
                text=True,
                timeout=60  # 60 second timeout
            )

            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")
                return {
                    "content": [{
                        "text": f"**PDF conversion failed**\n\n{result.stderr}"
                    }],
                    "status": "error"
                }

            pdf_path = os.path.join(temp_dir, document_filename.replace('.docx', '.pdf'))

            if not os.path.exists(pdf_path):
                return {
                    "content": [{
                        "text": "**PDF conversion failed**: Output file not created"
                    }],
                    "status": "error"
                }

            # Get total page count using pdfinfo (much faster than converting all pages)
            pdfinfo_result = subprocess.run(
                ['pdfinfo', pdf_path],
                capture_output=True,
                text=True,
                timeout=10
            )
            total_pages = 1  # default
            for line in pdfinfo_result.stdout.split('\n'):
                if line.startswith('Pages:'):
                    total_pages = int(line.split(':')[1].strip())
                    break

            # Validate all page numbers
            invalid_pages = [p for p in page_numbers if p > total_pages]
            if invalid_pages:
                return {
                    "content": [{
                        "text": f"**Invalid page(s): {invalid_pages}**\n\n"
                               f"Document has {total_pages} page(s)."
                    }],
                    "status": "error"
                }

            # Convert requested pages to images
            import io
            content = [{
                "text": f"**{document_filename}** - {len(page_numbers)} page(s) of {total_pages} total"
            }]

            for page_num in sorted(set(page_numbers)):  # Remove duplicates and sort
                logger.info(f"Converting page {page_num} to image...")
                images = convert_from_path(
                    pdf_path,
                    first_page=page_num,
                    last_page=page_num,
                    dpi=150
                )

                if images:
                    img_buffer = io.BytesIO()
                    images[0].save(img_buffer, format='PNG')
                    img_bytes = img_buffer.getvalue()

                    content.append({"text": f"**Page {page_num}:**"})
                    content.append({
                        "image": {
                            "format": "png",
                            "source": {"bytes": img_bytes}
                        }
                    })

            logger.info(f"Successfully generated {len(page_numbers)} preview(s)")

            text_blocks = [b for b in content if "text" in b]
            image_blocks = [b for b in content if "image" in b]
            return build_image_response(text_blocks, image_blocks, {
                "filename": document_filename,
                "page_numbers": sorted(set(page_numbers)),
                "total_pages": total_pages,
                "tool_type": "word_document",
                "user_id": user_id,
                "session_id": session_id,
                "hideImageInChat": True
            })

    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out")
        return {
            "content": [{
                "text": "**Conversion timed out**\n\nThe document may be too large or complex."
            }],
            "status": "error"
        }
    except Exception as e:
        logger.error(f"preview_word_page failed: {e}")
        return {
            "content": [{
                "text": f"**Failed to generate preview**\n\n{str(e)}"
            }],
            "status": "error"
        }


# --- Skill registration ---
register_skill("word-documents", tools=[create_word_document, modify_word_document, list_my_word_documents, read_word_document, preview_word_page])
